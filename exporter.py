import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# Color Theme Palettes for PDF
THEMES = {
    "Slate Modern": {
        "primary": colors.HexColor("#1e293b"),     # Slate 800
        "secondary": colors.HexColor("#475569"),   # Slate 600
        "text": colors.HexColor("#0f172a"),        # Slate 900
        "accent": colors.HexColor("#64748b")       # Slate 500
    },
    "Executive Navy": {
        "primary": colors.HexColor("#1e3a8a"),     # Navy 900
        "secondary": colors.HexColor("#2563eb"),   # Blue 600
        "text": colors.HexColor("#1e293b"),        # Slate 800
        "accent": colors.HexColor("#3b82f6")       # Blue 500
    },
    "Classic Charcoal": {
        "primary": colors.HexColor("#111827"),     # Gray 900
        "secondary": colors.HexColor("#374151"),   # Gray 700
        "text": colors.HexColor("#1f2937"),        # Gray 800
        "accent": colors.HexColor("#4b5563")       # Gray 600
    },
    "Emerald Professional": {
        "primary": colors.HexColor("#064e3b"),     # Emerald 900
        "secondary": colors.HexColor("#047857"),   # Emerald 700
        "text": colors.HexColor("#0f172a"),        # Slate 900
        "accent": colors.HexColor("#10b981")       # Emerald 500
    }
}

# --- WORD GENERATOR ---

def export_to_docx(resume_data) -> io.BytesIO:
    """
    Generates a beautifully styled Word Document (.docx) from structured resume data.
    """
    # If resume_data is a Pydantic model, convert to dict
    if hasattr(resume_data, "dict"):
        data = resume_data.dict()
    elif hasattr(resume_data, "model_dump"):
        data = resume_data.model_dump()
    else:
        data = resume_data

    doc = Document()
    
    # Page setup - Standard 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style settings
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 1. Header (Name, Contact Info)
    contact = data.get("contact_info", {})
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(contact.get("full_name", ""))
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B) # Primary Slate
    
    # Sub-header contact info
    contact_parts = []
    if contact.get("email"): contact_parts.append(contact["email"])
    if contact.get("phone"): contact_parts.append(contact["phone"])
    if contact.get("location"): contact_parts.append(contact["location"])
    if contact.get("linkedin"): contact_parts.append(contact["linkedin"])
    if contact.get("portfolio_or_website"): contact_parts.append(contact["portfolio_or_website"])
    
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run("  |  ".join(contact_parts))
    contact_run.font.size = Pt(9.5)
    contact_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B) # Accent
    
    # Helper to add section headers
    def add_section_header(title):
        header_para = doc.add_paragraph()
        header_para.paragraph_format.space_before = Pt(14)
        header_para.paragraph_format.space_after = Pt(4)
        header_run = header_para.add_run(title.upper())
        header_run.font.size = Pt(12.5)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Executive Blue
        
        # Add a subtle bottom border effect using drawing/line (we will use a paragraph bottom border or dash line)
        # Standard docx bottom border is XML-based. For simplicity and portability, we can add a simple horizontal rule
        border_para = doc.add_paragraph()
        border_para.paragraph_format.space_after = Pt(8)
        border_run = border_para.add_run("―" * 60)
        border_run.font.size = Pt(6)
        border_run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)

    # 2. Professional Summary
    if data.get("summary"):
        add_section_header("Professional Summary")
        summary_para = doc.add_paragraph(data["summary"])
        summary_para.paragraph_format.space_after = Pt(10)
        summary_para.paragraph_format.line_spacing = 1.15

    # 3. Work Experience
    if data.get("experience"):
        add_section_header("Professional Experience")
        for exp in data["experience"]:
            exp_title_para = doc.add_paragraph()
            exp_title_para.paragraph_format.space_before = Pt(6)
            exp_title_para.paragraph_format.space_after = Pt(2)
            
            # Job Title
            title_run = exp_title_para.add_run(exp.get("job_title", ""))
            title_run.font.bold = True
            title_run.font.size = Pt(11.5)
            
            # Company
            comp_run = exp_title_para.add_run(f" | {exp.get('company', '')}")
            comp_run.font.bold = True
            comp_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            
            # Location (Right align can be simulated or we just put it in line)
            loc_date_str = f" ({exp.get('location', '')}) — {exp.get('start_date', '')} to {exp.get('end_date', '')}"
            date_run = exp_title_para.add_run(loc_date_str)
            date_run.font.italic = True
            date_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

            # Bullets
            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(bullet, style='List Bullet')
                bp.paragraph_format.space_after = Pt(3)
                bp.paragraph_format.line_spacing = 1.15

    # 4. Skills
    if data.get("skills"):
        add_section_header("Core Skills")
        skills_para = doc.add_paragraph()
        skills_para.paragraph_format.space_after = Pt(10)
        skills_run = skills_para.add_run(", ".join(data["skills"]))
        skills_run.font.size = Pt(10.5)

    # 5. Projects
    if data.get("projects"):
        add_section_header("Projects")
        for proj in data["projects"]:
            proj_para = doc.add_paragraph()
            proj_para.paragraph_format.space_before = Pt(4)
            proj_para.paragraph_format.space_after = Pt(2)
            
            proj_title = proj.get("title", "")
            if proj.get("link"):
                proj_title += f" ({proj.get('link')})"
            
            title_run = proj_para.add_run(proj_title)
            title_run.font.bold = True
            
            tech_str = f" [Technologies: {', '.join(proj.get('technologies', []))}]"
            tech_run = proj_para.add_run(tech_str)
            tech_run.font.italic = True
            tech_run.font.size = Pt(9.5)
            tech_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

            desc_para = doc.add_paragraph(proj.get("description", ""))
            desc_para.paragraph_format.space_after = Pt(6)
            desc_para.paragraph_format.left_indent = Inches(0.2)

    # 6. Education
    if data.get("education"):
        add_section_header("Education")
        for edu in data["education"]:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_after = Pt(4)
            
            deg_str = f"{edu.get('degree', '')} in {edu.get('major', '')}"
            deg_run = edu_para.add_run(deg_str)
            deg_run.font.bold = True
            
            inst_str = f" | {edu.get('institution', '')}, {edu.get('location', '')}"
            inst_run = edu_para.add_run(inst_str)
            
            date_str = f" (Graduated: {edu.get('graduation_date', '')})"
            date_run = edu_para.add_run(date_str)
            date_run.font.italic = True
            date_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # 7. Certifications
    if data.get("certifications"):
        add_section_header("Certifications")
        certs_para = doc.add_paragraph()
        certs_para.paragraph_format.space_after = Pt(10)
        certs_run = certs_para.add_run(", ".join(data["certifications"]))
        certs_run.font.size = Pt(10.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- PDF GENERATOR (ReportLab) ---

def export_to_pdf(resume_data, theme_name: str = "Slate Modern") -> io.BytesIO:
    """
    Generates a high-quality, beautifully formatted PDF document using ReportLab.
    """
    # If resume_data is a Pydantic model, convert to dict
    if hasattr(resume_data, "dict"):
        data = resume_data.dict()
    elif hasattr(resume_data, "model_dump"):
        data = resume_data.model_dump()
    else:
        data = resume_data

    theme = THEMES.get(theme_name, THEMES["Slate Modern"])
    
    buffer = io.BytesIO()
    # 0.75 in (54 pt) margins all sides
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()
    
    # Custom styles
    style_name = ParagraphStyle(
        'ResumeName',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=22,
        leading=26,
        alignment=1, # Center
        textColor=theme["primary"]
    )
    
    style_contact = ParagraphStyle(
        'ResumeContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        alignment=1, # Center
        textColor=theme["accent"]
    )
    
    style_sec_heading = ParagraphStyle(
        'ResumeSectionHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=theme["primary"],
        spaceBefore=12,
        spaceAfter=3
    )
    
    style_summary = ParagraphStyle(
        'ResumeSummary',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=theme["text"],
        spaceAfter=8
    )

    style_item_title = ParagraphStyle(
        'ResumeItemTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=14,
        textColor=theme["text"]
    )
    
    style_item_sub = ParagraphStyle(
        'ResumeItemSub',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=9.5,
        leading=13,
        textColor=theme["secondary"]
    )

    style_bullet_text = ParagraphStyle(
        'ResumeBulletText',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=theme["text"]
    )
    
    story = []

    # 1. Header Section
    contact = data.get("contact_info", {})
    story.append(Paragraph(contact.get("full_name", ""), style_name))
    story.append(Spacer(1, 4))
    
    contact_parts = []
    if contact.get("email"): contact_parts.append(contact["email"])
    if contact.get("phone"): contact_parts.append(contact["phone"])
    if contact.get("location"): contact_parts.append(contact["location"])
    if contact.get("linkedin"): contact_parts.append(contact["linkedin"])
    if contact.get("portfolio_or_website"): contact_parts.append(contact["portfolio_or_website"])
    
    contact_text = "  •  ".join(contact_parts)
    story.append(Paragraph(contact_text, style_contact))
    story.append(Spacer(1, 10))

    # Helper function to generate section headings with bottom border lines
    def build_section_header(title):
        header_para = Paragraph(title.upper(), style_sec_heading)
        # Using a table with bottom border to create a clean divider line
        divider_table = Table([[""]], colWidths=[504], rowHeights=[2])
        divider_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), theme["primary"]),
            ('TOPPADDING', (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ]))
        return [header_para, divider_table, Spacer(1, 6)]

    # 2. Professional Summary
    if data.get("summary"):
        story.extend(build_section_header("Professional Summary"))
        story.append(Paragraph(data["summary"], style_summary))
        story.append(Spacer(1, 6))

    # 3. Work Experience
    if data.get("experience"):
        story.extend(build_section_header("Professional Experience"))
        for exp in data["experience"]:
            # Title & Dates side-by-side using Table
            job_title_html = f"<b>{exp.get('job_title', '')}</b>"
            company_html = f" &nbsp;|&nbsp; <b>{exp.get('company', '')}</b>"
            left_p = Paragraph(f"{job_title_html}{company_html}", style_item_title)
            
            loc_date_html = f"{exp.get('location', '')} &nbsp;•&nbsp; {exp.get('start_date', '')} - {exp.get('end_date', '')}"
            right_p = Paragraph(f"<font color='{theme['accent'].hexval()}'>{loc_date_html}</font>", ParagraphStyle(
                'RightAlign', parent=style_item_sub, alignment=2 # Right
            ))
            
            exp_table = Table([[left_p, right_p]], colWidths=[320, 184])
            exp_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 2),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ]))
            story.append(exp_table)
            
            # Bullets
            for bullet in exp.get("bullets", []):
                bullet_bullet = Paragraph("•", ParagraphStyle('BulletBul', parent=style_bullet_text, alignment=1))
                bullet_para = Paragraph(bullet, style_bullet_text)
                
                bullet_table = Table([[bullet_bullet, bullet_para]], colWidths=[15, 489])
                bullet_table.setStyle(TableStyle([
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('LEFTPADDING', (0,0), (-1,-1), 0),
                    ('RIGHTPADDING', (0,0), (-1,-1), 0),
                    ('TOPPADDING', (0,0), (-1,-1), 1),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ]))
                story.append(bullet_table)
            
            story.append(Spacer(1, 6))

    # 4. Skills
    if data.get("skills"):
        story.extend(build_section_header("Skills"))
        skills_text = ", ".join(data["skills"])
        story.append(Paragraph(skills_text, style_bullet_text))
        story.append(Spacer(1, 8))

    # 5. Projects
    if data.get("projects"):
        story.extend(build_section_header("Projects"))
        for proj in data["projects"]:
            # Project Title & Tech Stack side-by-side
            proj_title = f"<b>{proj.get('title', '')}</b>"
            if proj.get('link'):
                proj_title += f" &nbsp;(<font color='{theme['secondary'].hexval()}'><u>{proj.get('link')}</u></font>)"
            left_p = Paragraph(proj_title, style_item_title)
            
            tech_str = f"Technologies: {', '.join(proj.get('technologies', []))}"
            right_p = Paragraph(f"<font color='{theme['accent'].hexval()}'>{tech_str}</font>", ParagraphStyle(
                'RightAlignProj', parent=style_item_sub, alignment=2 # Right
            ))
            
            proj_table = Table([[left_p, right_p]], colWidths=[240, 264])
            proj_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 2),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ]))
            story.append(proj_table)
            
            # Project Description
            desc_bullet = Paragraph("•", ParagraphStyle('BulletBulProj', parent=style_bullet_text, alignment=1))
            desc_para = Paragraph(proj.get("description", ""), style_bullet_text)
            desc_table = Table([[desc_bullet, desc_para]], colWidths=[15, 489])
            desc_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
            ]))
            story.append(desc_table)
            story.append(Spacer(1, 6))

    # 6. Education
    if data.get("education"):
        story.extend(build_section_header("Education"))
        for edu in data["education"]:
            major_deg = f"<b>{edu.get('degree', '')} in {edu.get('major', '')}</b>"
            left_p = Paragraph(f"{major_deg} &nbsp;|&nbsp; {edu.get('institution', '')}, {edu.get('location', '')}", style_bullet_text)
            
            date_p = Paragraph(f"<font color='{theme['accent'].hexval()}'>Graduated: {edu.get('graduation_date', '')}</font>", ParagraphStyle(
                'RightAlignEdu', parent=style_item_sub, alignment=2 # Right
            ))
            
            edu_table = Table([[left_p, date_p]], colWidths=[380, 124])
            edu_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 2),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ]))
            story.append(edu_table)
            story.append(Spacer(1, 4))
        story.append(Spacer(1, 4))

    # 7. Certifications
    if data.get("certifications"):
        story.extend(build_section_header("Certifications & Licenses"))
        certs_text = ", ".join(data["certifications"])
        story.append(Paragraph(certs_text, style_bullet_text))

    # Build the document
    doc.build(story)
    buffer.seek(0)
    return buffer
