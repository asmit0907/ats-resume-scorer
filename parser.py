import docx
from pypdf import PdfReader
import io

def parse_pdf(file_stream: io.BytesIO) -> str:
    """
    Extracts text from an uploaded PDF file stream.
    """
    try:
        reader = PdfReader(file_stream)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")

def parse_docx(file_stream: io.BytesIO) -> str:
    """
    Extracts text from an uploaded DOCX file stream.
    """
    try:
        doc = docx.Document(file_stream)
        text = []
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text:
                text.append(para.text)
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Clean cell text duplicates
                    cell_text = cell.text.strip()
                    if cell_text and (not text or text[-1] != cell_text):
                        text.append(cell_text)
        return "\n".join(text).strip()
    except Exception as e:
        raise ValueError(f"Failed to parse Word Document: {str(e)}")

def extract_text(file_stream: io.BytesIO, file_name: str) -> str:
    """
    Determines file type and extracts text from it.
    """
    if file_name.lower().endswith(".pdf"):
        return parse_pdf(file_stream)
    elif file_name.lower().endswith((".docx", ".doc")):
        # Note: python-docx only supports docx, doc is binary format and needs special handle,
        # but modern users use docx. We can warn them if .doc fails.
        return parse_docx(file_stream)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
